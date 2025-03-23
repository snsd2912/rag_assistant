import asyncio
import os
import multiprocessing
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents  import Document
from docx import Document as DocxDocument
from langchain_text_splitters import CharacterTextSplitter
from dotenv import load_dotenv
from pathlib import Path
from fastapi import APIRouter, HTTPException, WebSocket, WebSocketDisconnect,File, UploadFile
from langdetect import detect
from models.question import Question
import numpy as np
from langchain_postgres import PGVector
from langchain_postgres.vectorstores import PGVector
import os
from langchain_openai import ChatOpenAI, OpenAIEmbeddings

CONNECTION_STRING = os.getenv("VECTOR_DB_URL")

# Initialize FastAPI app
router = APIRouter()

# Load environment variables
load_dotenv()

API_KEY = os.getenv("API_KEY")

# Initialize chat model (GPT-4o)
chat_model = ChatOpenAI(
    model="gpt-3.5-turbo",
    openai_api_key=API_KEY
)

# Initialize embeddings
embeddings = OpenAIEmbeddings(
    model="text-embedding-ada-002",
    openai_api_key=API_KEY
)

# Global vector_store
vector_store = None

UPLOAD_DIR = Path("./data_warehouse")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

def load_docx(file_path):
    """
    Custom function to load content from a .docx file and return it as LangChain Document objects.
    """
    doc = DocxDocument(file_path)
    content = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Ignore empty paragraphs
            content.append(paragraph.text.strip())

    # Return as a list of LangChain Document objects
    return [Document(page_content="\n".join(content), metadata={"source": str(file_path)})]

def load_and_embed_pdfs_and_docs(folder):
    """
    Load PDFs and DOCX files from a folder and create embeddings.
    """
    docs = []
    pdf_files = Path(folder).glob("*.pdf")
    docx_files = Path(folder).glob("*.docx")

    for pdf_file in pdf_files:
        loader = PyPDFLoader(str(pdf_file))
        documents = loader.load()
        docs.extend(documents)

    for docx_file in docx_files:
        documents = load_docx(str(docx_file))
        docs.extend(documents)

    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)

    split_docs = text_splitter.split_documents(docs)

    global vector_store
    if vector_store is None:
        vector_store = PGVector.from_documents(
            documents=split_docs,
            embedding=embeddings,
            collection_name="vectorstore",
            connection=CONNECTION_STRING,
            use_jsonb=True,
        )
        # vector_store = FAISS.from_documents(split_docs, embeddings)
    else:
        vector_store.add_documents(split_docs)

    return vector_store

def find_closest_pdf(prompt, vector_store):
    """
    Search for the closest PDF to a user's prompt.
    """
    result = vector_store.similarity_search(prompt, k=1)
    return result[0] if result else None

def summarize_pdf_with_gpt4o(text, detected_language):
    """
    Summarize a PDF's content using Azure OpenAI (gpt-4o-mini).
    """
    prompt = ChatPromptTemplate.from_template("""
            You are a helpful assistant. Please summarize the main points of the following document and translate to language {language} you can detect:

            {content}

            Provide a clear and concise summary as Bullet points with language {language}. Only get translate version.
        """)
    
    content_chain = prompt | chat_model

    response = content_chain.invoke({"content": text, "language": detected_language})
    
    return response.content

def process_user_prompt(prompt, pdf_folder, detected_language):
    """
    Full pipeline: find the closest PDF, extract content, and summarize.
    """
    # Load PDFs and create embeddings
    global vector_store

    # Find the closest PDF
    closest_pdf = find_closest_pdf(prompt, vector_store)

    if closest_pdf:
        pdf_content = closest_pdf.page_content
        summary = summarize_pdf_with_gpt4o(pdf_content, detected_language)
        return {
            "closest_pdf": closest_pdf.metadata['source'],
            "summary": summary
        }
    else:
        return {
            "error": "We still collecting the documents relate your question."
        }

@router.post("/ask")
async def ask_question(payload: Question):
    """
    API to process user question and return summary of the most relevant PDF.
    """
    pdf_folder = "./data_warehouse"
    question = payload.question

    if not question:
        raise HTTPException(status_code=400, detail="Question cannot be empty.")

    detected_language = detect(question)

    if detected_language != "vi":
        detected_language = "en"

    file_names = [file.name for file in UPLOAD_DIR.iterdir() if file.is_file()]

    prompt = ChatPromptTemplate.from_template("""
               Bạn là một người hỗ trợ tuyệt vời, tôi có một danh sách các file bên dưới:

               {folders}

               Hãy tìm cho tôi file có tên mà có khả năng xuất hiện trong câu hỏi. Đây là câu hỏi:
               
               {question}
               
               Nếu không có file có tên thích hợp câu trả lời của bạn chỉ là:No
               
           """)

    content_chain = prompt | chat_model
    response = content_chain.invoke({"folders": file_names, "question": question})
    print(f"response.content= {response.content}")
    if response.content == "No":
        if detected_language != "vi":
            return {"summary": "We still collecting data for your question."}
        return {"summary": "chúng tôi vẫn đang tập hợp giữ liệu cho câu hỏi của bạn."}
    result = process_user_prompt(question, pdf_folder, detected_language)

    if not result:
        raise HTTPException(status_code=404, detail="No relevant document found.")

    if "error" in result:
        return {
            "summary": result["error"]
        }

    if detected_language != "vi":
        answer = " You can refer steps:\n " + result["summary"]
    else:
        answer = " Bạn có thể tham khảo các bước sau:\n " + result["summary"]
    return {
        "closest_pdf": result["closest_pdf"],
        "summary": answer
    }

def update_vector_store_with_file(file_path):
    """
    Update the vector_store with a newly uploaded file.
    """
    global vector_store
    file_extension = file_path.suffix.lower()
    if file_extension == ".pdf":
        loader = PyPDFLoader(str(file_path))
        documents = loader.load()
    elif file_extension == ".docx":
        documents = load_docx(str(file_path))
    else:
        raise ValueError("Unsupported file type.")

    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    split_docs = text_splitter.split_documents(documents)

    if vector_store is None:
        vector_store = PGVector.from_documents(
            documents=split_docs,
            embedding=embeddings,
            collection_name="vectorstore",
            connection=CONNECTION_STRING,
            use_jsonb=True,
        )
        # vector_store = FAISS.from_documents(split_docs, embeddings)
    else:
        vector_store.add_documents(split_docs)

@router.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """
    Endpoint to upload a file and update the vector_store.
    """
    file_extension = file.filename.split('.')[-1]
    if file_extension not in ["pdf", "docx"]:
        return {"error": "File type not allowed. Only .pdf and .docx are supported."}

    file_path = UPLOAD_DIR / file.filename

    if file_path.exists():
        return {"error": f"File {file.filename} already exists."}

    with open(file_path, "wb") as f:
        f.write(await file.read())

    try:
        process = multiprocessing.Process(target=update_vector_store_with_file(file_path))
        process.start()
        update_vector_store_with_file(file_path)
    except Exception as e:
        return {"error": f"Failed to update vector store: {str(e)}"}

    return {"message": f"{file.filename} File uploaded successfully!"}


async def stream_lines(data):
    lines = data.split("\n")
    for line in lines:
        yield line
        await asyncio.sleep(1)

@router.on_event("startup")
async def initialize_vector_store():
    """
    Load the vector_store during app startup.
    """
    global vector_store
    folder = "./data_warehouse"
    vector_store = load_and_embed_pdfs_and_docs(folder)
    print("Vector store initialized!")
